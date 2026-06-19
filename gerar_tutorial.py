from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

AZUL      = RGBColor(0x00, 0x2B, 0x5C)
AZUL_MED  = RGBColor(0x00, 0x56, 0xA2)
VERDE     = RGBColor(0x00, 0x9A, 0x44)
AMARELO   = RGBColor(0xFF, 0xC2, 0x00)
CINZA     = RGBColor(0xF2, 0xF2, 0xF2)
BRANCO    = RGBColor(0xFF, 0xFF, 0xFF)
PRETO     = RGBColor(0x22, 0x22, 0x22)

doc = Document()

# Margens
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

def set_font(run, size=11, bold=False, italic=False, color=PRETO, name="Calibri"):
    run.font.name  = name
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.italic = italic
    run.font.color.rgb = color

def add_heading(doc, text, level=1, color=AZUL):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    sizes = {1: 20, 2: 15, 3: 13}
    set_font(run, size=sizes.get(level, 13), bold=True, color=color)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    return p

def add_body(doc, text, color=PRETO, italic=False, size=11, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, size=size, color=color, italic=italic)
    p.paragraph_format.space_after = Pt(space_after)
    return p

def add_bullet(doc, text, color=PRETO, indent=0):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    set_font(run, size=11, color=color)
    p.paragraph_format.left_indent = Inches(0.3 + indent * 0.3)
    p.paragraph_format.space_after = Pt(3)
    return p

def add_code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.4)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    set_font(run, size=11, bold=True, color=AZUL_MED, name="Courier New")
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), "E8F0FE")
    p._p.get_or_add_pPr().append(shading)
    return p

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        set_font(run, size=10, bold=True, color=BRANCO)
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "002B5C")
        tcPr.append(shd)
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        fill = "F2F2F2" if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            set_font(run, size=10, color=PRETO)
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), fill)
            tcPr.append(shd)
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(width)
    doc.add_paragraph()
    return table

def add_separator(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "002B5C")
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(8)

def add_destaque(doc, text, bg="FFF3CD"):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(8)
    run = p.add_run(text)
    set_font(run, size=11, italic=True, color=PRETO)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), bg)
    p._p.get_or_add_pPr().append(shading)

# ══════════════════════════════════════════════════════
# CAPA
# ══════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(40)
run = p.add_run("BRASILSEG")
set_font(run, size=13, bold=True, color=AZUL_MED)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(20)
run = p.add_run("Tutorial Prático")
set_font(run, size=14, color=PRETO)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Como usar o Claude Code no dia a dia\nda Gerente de Projetos")
set_font(run, size=26, bold=True, color=AZUL)
p.paragraph_format.space_after = Pt(10)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Skills GP: /gp-status  •  /gp-ata  •  /gp-standup  •  /gp-metricas  •  /gp-priorizacao")
set_font(run, size=12, italic=True, color=AZUL_MED)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(30)
run = p.add_run("Andressa Marquito  |  Gerente de Portfolio de Projetos\ndessajf@gmail.com")
set_font(run, size=11, color=PRETO)

doc.add_page_break()

# ══════════════════════════════════════════════════════
# SUMÁRIO MANUAL
# ══════════════════════════════════════════════════════
add_heading(doc, "Sumário", 1)
sumario = [
    "1. O que é o Claude Code e como ele funciona",
    "2. Como acessar e digitar os comandos",
    "3. Rotina diária da GP com o Claude",
    "4. Guia de cada skill — passo a passo",
    "   4.1  /gp-status  — Relatório de Status",
    "   4.2  /gp-ata     — Ata de Reunião",
    "   4.3  /gp-standup — Standup Diário",
    "   4.4  /gp-metricas — Métricas do Projeto",
    "   4.5  /gp-priorizacao — Priorização de Portfólio",
    "5. Exemplos reais de uso (BrasilSeg)",
    "6. Dicas avançadas e boas práticas",
    "7. Perguntas frequentes",
]
for item in sumario:
    add_body(doc, item, color=AZUL_MED, size=11, space_after=3)

doc.add_page_break()

# ══════════════════════════════════════════════════════
# 1. O QUE É O CLAUDE CODE
# ══════════════════════════════════════════════════════
add_heading(doc, "1. O que é o Claude Code e como ele funciona", 1)
add_separator(doc)

add_body(doc, "O Claude Code é um assistente de inteligência artificial desenvolvido pela Anthropic que roda diretamente no seu computador, dentro do ambiente de desenvolvimento (VS Code). Diferente de chatbots comuns, o Claude Code tem acesso ao seu projeto, pode criar arquivos, gerar documentos e executar tarefas complexas com um simples comando.")

add_heading(doc, "O que ele faz pela GP:", 2, color=AZUL_MED)
add_bullet(doc, "Gera relatórios de status completos em segundos")
add_bullet(doc, "Estrutura atas de reunião no padrão profissional")
add_bullet(doc, "Conduz e registra standups diários")
add_bullet(doc, "Calcula métricas (CPI, SPI, Velocity, Lead Time) automaticamente")
add_bullet(doc, "Prioriza portfólio de produtos com matriz objetiva")
add_bullet(doc, "Cria apresentações em PowerPoint e documentos Word")
add_bullet(doc, "Responde perguntas, sugere soluções e redige comunicados")

add_destaque(doc, "💡 Pense no Claude como uma analista sênior disponível 24h, que conhece suas metodologias, seus projetos e nunca esquece o que você pediu na última conversa.")

doc.add_page_break()

# ══════════════════════════════════════════════════════
# 2. COMO ACESSAR
# ══════════════════════════════════════════════════════
add_heading(doc, "2. Como acessar e digitar os comandos", 1)
add_separator(doc)

add_heading(doc, "Abrindo o Claude Code", 2, color=AZUL_MED)
add_bullet(doc, "Abra o VS Code no seu computador")
add_bullet(doc, "No painel lateral esquerdo, clique no ícone do Claude (ou pressione Ctrl+Shift+P e busque 'Claude')")
add_bullet(doc, "A janela de chat abrirá na lateral direita")

add_heading(doc, "Como digitar um comando (skill)", 2, color=AZUL_MED)
add_body(doc, "Basta digitar o nome do comando com a barra / no início e pressionar Enter:")
add_code_block(doc, "/gp-status")
add_body(doc, "O Claude vai responder imediatamente e fazer as perguntas necessárias para completar a tarefa.")

add_heading(doc, "Comandos disponíveis", 2, color=AZUL_MED)
add_table(doc,
    ["Comando", "Quando usar", "Tempo médio"],
    [
        ["/gp-status",      "Toda sexta-feira ou antes de reunião de acompanhamento", "3 minutos"],
        ["/gp-ata",         "Logo após qualquer reunião",                              "5 minutos"],
        ["/gp-standup",     "Todo dia útil pela manhã",                               "2 minutos"],
        ["/gp-metricas",    "Quinzenalmente ou antes de comitê",                      "4 minutos"],
        ["/gp-priorizacao", "No início de cada ciclo de planejamento",                "5 minutos"],
    ],
    col_widths=[1.8, 3.5, 1.5]
)

add_destaque(doc, "✅ Dica: você não precisa saber programação. Os comandos funcionam como uma conversa normal — o Claude pergunta, você responde, ele entrega.")

doc.add_page_break()

# ══════════════════════════════════════════════════════
# 3. ROTINA DIÁRIA
# ══════════════════════════════════════════════════════
add_heading(doc, "3. Rotina diária da GP com o Claude", 1)
add_separator(doc)

add_table(doc,
    ["Horário", "Atividade", "Comando Claude", "Entregável"],
    [
        ["08h00", "Standup da equipe",              "/gp-standup",     "Registro do standup + impedimentos"],
        ["08h30", "Acompanhamento de tarefas",      "Conversa livre",  "Respostas, análises e sugestões"],
        ["12h00", "Pós-reunião de alinhamento",     "/gp-ata",         "Ata formatada e pronta para envio"],
        ["16h00", "Atualização de status",          "/gp-status",      "Relatório RAG atualizado"],
        ["Sex 17h", "Fechamento semanal / métricas", "/gp-metricas",   "Painel de indicadores da semana"],
        ["Quinzenal", "Revisão de portfólio",       "/gp-priorizacao", "Ranking atualizado de projetos"],
    ],
    col_widths=[1.2, 2.3, 1.8, 2.5]
)

add_body(doc, "Além dos comandos, você pode conversar livremente com o Claude a qualquer momento:")
add_bullet(doc, "\"Me ajude a redigir um e-mail cobrando entrega do fornecedor X\"")
add_bullet(doc, "\"Qual a diferença entre Scrum e Kanban para um time de seguros?\"")
add_bullet(doc, "\"Crie uma apresentação sobre o projeto Vida Rural para a diretoria\"")
add_bullet(doc, "\"Analise esses dados e me diga se o projeto está em risco\"")

doc.add_page_break()

# ══════════════════════════════════════════════════════
# 4. GUIA DE CADA SKILL
# ══════════════════════════════════════════════════════
add_heading(doc, "4. Guia de cada skill — passo a passo", 1)
add_separator(doc)

# 4.1 gp-status
add_heading(doc, "4.1  /gp-status — Relatório de Status", 2, color=VERDE)
add_body(doc, "Gera um relatório de status completo com indicadores RAG (🟢🟡🔴) pronto para enviar à liderança.")
add_heading(doc, "Como usar:", 3, color=PRETO)
add_bullet(doc, "Passo 1: Digite /gp-status e pressione Enter")
add_bullet(doc, "Passo 2: O Claude vai perguntar: nome do projeto, período, o que foi concluído, impedimentos e próximas entregas")
add_bullet(doc, "Passo 3: Responda cada pergunta normalmente, como em uma conversa")
add_bullet(doc, "Passo 4: O Claude entrega o relatório formatado, com tabelas e RAG")
add_heading(doc, "Exemplo de uso real (BrasilSeg):", 3, color=PRETO)
add_destaque(doc, 'Você digita: /gp-status\n\nClaude pergunta: "Qual o nome do projeto?"\nVocê responde: "Lançamento Seguro Rural Safra 2026 — Ramo Rural"\n\nClaude pergunta: "O que foi concluído esta semana?"\nVocê responde: "Concluímos a análise de viabilidade e enviamos para o comitê técnico."\n\n→ Claude entrega o relatório completo com RAG, tabela de riscos e próximas entregas.', bg="E8F5E9")

# 4.2 gp-ata
add_heading(doc, "4.2  /gp-ata — Ata de Reunião", 2, color=VERDE)
add_body(doc, "Estrutura qualquer ata no padrão: Pauta → Discussões → Decisões → Ações com responsáveis e prazos.")
add_heading(doc, "Como usar:", 3, color=PRETO)
add_bullet(doc, "Passo 1: Assim que a reunião terminar, abra o Claude e digite /gp-ata")
add_bullet(doc, "Passo 2: Informe: tipo de reunião, participantes e data")
add_bullet(doc, "Passo 3: Narre os pontos discutidos de forma livre — pode ser resumido")
add_bullet(doc, "Passo 4: O Claude organiza tudo no formato correto com ações e responsáveis")
add_heading(doc, "Exemplo de uso real (BrasilSeg):", 3, color=PRETO)
add_destaque(doc, 'Você digita: /gp-ata\n\nClaude pergunta o tipo: "Planning Sprint 12 — Squad Danos"\nParticipantes: "Andressa, João (TI), Maria (Produtos), Carlos (Compliance)"\n\nVocê narra: "Discutimos o backlog do produto, João disse que a integração com o legado vai atrasar 1 semana, Carlos pediu revisão do texto da apólice antes do lançamento."\n\n→ Claude entrega ata completa com decisões em negrito e tabela de ações com prazos.', bg="E8F5E9")

# 4.3 gp-standup
add_heading(doc, "4.3  /gp-standup — Standup Diário", 2, color=VERDE)
add_body(doc, "Conduz o standup das 3 perguntas clássicas e registra impedimentos automaticamente.")
add_heading(doc, "Como usar:", 3, color=PRETO)
add_bullet(doc, "Passo 1: Digite /gp-standup às 8h antes da reunião")
add_bullet(doc, "Passo 2: Informe o nome do projeto e os membros presentes")
add_bullet(doc, "Passo 3: Para cada membro, informe o que fez ontem, o que fará hoje e impedimentos")
add_bullet(doc, "Passo 4: O Claude gera o registro e destaca impedimentos para ação imediata")
add_destaque(doc, "💡 Dica: você pode informar todos os membros de uma vez em formato livre — o Claude organiza automaticamente.", bg="FFF8E1")

# 4.4 gp-metricas
add_heading(doc, "4.4  /gp-metricas — Métricas do Projeto", 2, color=VERDE)
add_body(doc, "Calcula CPI, SPI, Velocity, Lead Time e gera painel de indicadores com análise e recomendações.")
add_heading(doc, "Como usar:", 3, color=PRETO)
add_bullet(doc, "Passo 1: Digite /gp-metricas")
add_bullet(doc, "Passo 2: Escolha quais métricas calcular (ou 'todas')")
add_bullet(doc, "Passo 3: Forneça os dados (ex: Valor Agregado = R$80.000 / Custo Real = R$90.000)")
add_bullet(doc, "Passo 4: O Claude calcula, interpreta e recomenda ações")
add_heading(doc, "Referência rápida de interpretação:", 3, color=PRETO)
add_table(doc,
    ["Métrica", "Resultado", "O que significa", "Ação"],
    [
        ["CPI", "< 1,0", "🔴 Acima do orçamento", "Revisar custos imediatamente"],
        ["CPI", "= 1,0", "🟢 No orçamento",       "Manter"],
        ["CPI", "> 1,0", "🟢 Abaixo do orçamento","Excelente — monitorar"],
        ["SPI", "< 1,0", "🔴 Atrasado",           "Replanejar entregas"],
        ["SPI", "= 1,0", "🟢 No prazo",           "Manter"],
        ["SPI", "> 1,0", "🟢 Adiantado",          "Possível antecipação"],
    ],
    col_widths=[1.2, 1.2, 2.5, 2.5]
)

# 4.5 gp-priorizacao
add_heading(doc, "4.5  /gp-priorizacao — Priorização de Portfólio", 2, color=VERDE)
add_body(doc, "Aplica a Matriz de Priorização do Funil Ágil BrasilSeg e gera ranking objetivo dos projetos.")
add_heading(doc, "Como usar:", 3, color=PRETO)
add_bullet(doc, "Passo 1: Digite /gp-priorizacao no início do ciclo de planejamento")
add_bullet(doc, "Passo 2: Liste os projetos de cada ramo (Danos, Vida/Prestamista, Rural)")
add_bullet(doc, "Passo 3: Para cada projeto, informe: receita esperada, alinhamento estratégico, complexidade regulatória e esforço")
add_bullet(doc, "Passo 4: O Claude calcula o score ponderado e entrega o ranking final")
add_destaque(doc, "⭐ Esse comando é especialmente útil antes de reuniões de diretoria — você chega com dados, não com opinião.", bg="E3F2FD")

doc.add_page_break()

# ══════════════════════════════════════════════════════
# 5. EXEMPLOS REAIS
# ══════════════════════════════════════════════════════
add_heading(doc, "5. Exemplos reais de uso (BrasilSeg)", 1)
add_separator(doc)

cenarios = [
    ("Cenário 1: Segunda de manhã — Planning da Sprint",
     "A equipe do ramo Rural vai iniciar a Sprint 8. Você precisa conduzir a planning e registrar o backlog priorizado.",
     ["/gp-standup — para registrar o início da sprint e as entregas planejadas",
      "Conversa livre: 'Crie um template de backlog para a Sprint 8 do projeto Seguro Rural Safra 2026'",
      "O Claude entrega backlog formatado com histórias, critérios de aceite e estimativa"]),
    ("Cenário 2: Quarta à tarde — Reunião com a diretoria",
     "Você precisa apresentar o status do portfólio dos 3 ramos para a diretoria em 30 minutos.",
     ["/gp-status — execute para cada projeto principal",
      "/gp-priorizacao — para mostrar o ranking atualizado do portfólio",
      "Conversa livre: 'Crie um sumário executivo de 1 página com os status dos 3 ramos'"]),
    ("Cenário 3: Sexta — Fechamento semanal",
     "Fim de semana. Você precisa enviar o relatório consolidado para a liderança.",
     ["/gp-metricas — calcule CPI e SPI dos projetos em andamento",
      "/gp-status — gere o relatório semanal",
      "Conversa livre: 'Redija um e-mail executivo com o status da semana para enviar à diretoria'"]),
    ("Cenário 4: Novo produto surgiu — precisa priorizar",
     "A área comercial do ramo Danos trouxe uma ideia de novo produto. Como encaixar no portfólio?",
     ["/gp-priorizacao — inclua o novo produto e veja onde ele fica no ranking",
      "Conversa livre: 'Faça um Lean Canvas básico para o produto X do ramo Danos'",
      "O Claude estrutura a análise inicial da Fase 1 do Funil Ágil"]),
]

for titulo, descricao, passos in cenarios:
    add_heading(doc, titulo, 2, color=AZUL_MED)
    add_body(doc, descricao, italic=True)
    add_body(doc, "Como usar o Claude:", color=VERDE, bold=False)
    for passo in passos:
        add_bullet(doc, passo)
    doc.add_paragraph()

doc.add_page_break()

# ══════════════════════════════════════════════════════
# 6. DICAS AVANÇADAS
# ══════════════════════════════════════════════════════
add_heading(doc, "6. Dicas avançadas e boas práticas", 1)
add_separator(doc)

dicas = [
    ("Contextualize sempre", "Antes de pedir qualquer coisa, diga ao Claude o contexto: 'Sou GP do portfólio BrasilSeg, trabalhamos com os ramos Danos, Vida/Prestamista e Rural. Preciso de...' — ele responderá de forma muito mais precisa."),
    ("Combine skills em sequência", "Você pode usar /gp-metricas e depois pedir: 'Inclua essas métricas no relatório de status que acabei de gerar.' O Claude mantém o contexto da conversa."),
    ("Peça ajustes", "Se o resultado não ficou exatamente como você queria, diga: 'Mude o tom para mais executivo' ou 'Adicione uma coluna de prazo na tabela de riscos.' O Claude corrige imediatamente."),
    ("Salve os resultados", "Após gerar um relatório ou ata, você pode pedir: 'Salve isso em um arquivo Word na minha Área de Trabalho.' O Claude cria o arquivo automaticamente."),
    ("Use para redigir comunicados", "Além dos comandos, o Claude é excelente para redigir: e-mails de cobrança, comunicados de mudança, apresentações executivas, descrições de risco e planos de ação."),
    ("Peça análises comparativas", "Exemplo: 'Compare o desempenho do ramo Danos com o Rural nos últimos 2 meses com base nesses dados...' — O Claude analisa e sugere ações."),
]

for titulo, texto in dicas:
    add_heading(doc, f"► {titulo}", 3, color=AZUL_MED)
    add_body(doc, texto)

doc.add_page_break()

# ══════════════════════════════════════════════════════
# 7. FAQ
# ══════════════════════════════════════════════════════
add_heading(doc, "7. Perguntas frequentes", 1)
add_separator(doc)

faqs = [
    ("O Claude salva minhas conversas?",
     "Dentro da mesma sessão, o Claude lembra de tudo que foi dito. Entre sessões diferentes, ele mantém memórias configuradas (como seu perfil de GP). Dados sensíveis de projetos não devem ser compartilhados em texto aberto."),
    ("Posso usar o Claude para projetos confidenciais?",
     "O Claude roda localmente no seu computador via VS Code. Consulte a política de segurança da BrasilSeg sobre uso de IA antes de inserir dados sensíveis de clientes ou produtos em aprovação."),
    ("O que faço se o comando não funcionar?",
     "Verifique se digitou corretamente com a barra: /gp-status. Se o problema persistir, feche e reabra o VS Code. Os comandos ficam em C:\\Users\\dessa\\.claude\\commands\\"),
    ("Posso criar novos comandos personalizados?",
     "Sim! Novos comandos podem ser criados a qualquer momento. Basta pedir ao Claude: 'Crie um novo comando /gp-risco para gerenciar a matriz de riscos do projeto.'"),
    ("O Claude pode acessar meu e-mail ou calendário?",
     "Não de forma automática. Mas você pode copiar o conteúdo de um e-mail ou reunião e pedir ao Claude que analise, responda ou estruture."),
    ("Posso usar o Claude no celular?",
     "O Claude Code (VS Code) é para desktop. Para uso mobile, acesse claude.ai pelo navegador do celular — os skills personalizados não estarão disponíveis, mas você pode conversar normalmente."),
]

for pergunta, resposta in faqs:
    add_heading(doc, f"P: {pergunta}", 3, color=AZUL_MED)
    add_body(doc, f"R: {resposta}")
    doc.add_paragraph()

# Rodapé final
add_separator(doc)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Andressa Marquito  •  Gerente de Portfolio de Projetos  •  BrasilSeg\ndessajf@gmail.com  •  Tutorial gerado com Claude Code")
set_font(run, size=10, italic=True, color=AZUL_MED)

output = r"c:\Users\dessa\OneDrive\Área de Trabalho\Tutorial_Claude_GP_BrasilSeg.docx"
doc.save(output)
print(f"Arquivo salvo: {output}")
